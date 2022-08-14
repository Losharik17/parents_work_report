import telebot
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from num2t4ru import num2text
# from docx2pdf import convert
import re
import datetime
import subprocess


bot = telebot.TeleBot('5300733750:AAHFpMDYrmWNWopow41US1pREfHptqTPJ_E')

data = []
temp_date = None
last_date = []
first_str = ['№', 'Товар (Работы, услуги)', 'Кол-во', 'Ед.', 'Цена', 'Сумма']
month = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля', 'августа', 'сентября',
         'октября', 'ноября', 'декабря']
price = 2000
sum_hours = 0
total_sum = 0
number = 1031278


@bot.message_handler(content_types=['text'])
def get_text_messages(message):
    if message.text == "/doc":
        bot.send_message(message.from_user.id, "Начинаем формирование документа")
        bot.send_message(message.from_user.id, "Укажи дату в формате дд или дд.мм или дд.мм.гг ")
        bot.register_next_step_handler(message, get_date)
    elif message.text == "/price":
        bot.send_message(message.from_user.id, "Введи новую цену услуги")
        bot.register_next_step_handler(message, change_price)
    elif message.text == "/help":
        bot.send_message(message.from_user.id, "Нажми /doc, чтобы создать новый документ\n\n"
                                               "Нажми /price, чтобы изменить цену\n\n"
                                               "При вводе даты, если ты не укажешь месяц или год, то будет установлен "
                                               "последний указанный месяц/год.\n"
                                               "Если ранее не было ничего указано, то установится текущий месяц/год")
    else:
        bot.send_message(message.from_user.id, "Я тебя не понимаю. Напиши /help")

# xxx
def change_price(message):
    global price
    if re.match('\d\d\d\d', message.text) is not None:
        try:
            price = int(message.text)
        except:
            pass


def get_date(message):
    global temp_date, data, last_date

    if not last_date:
        now = datetime.datetime.now()
        last_date = [str(now.day), str(now.month), str(now.year)]

    temp_date = message.text.replace(',', '.')

    if temp_date == '/end':
        bot.send_message(message.from_user.id, 'Документ формируется')
        create_doc()
        bot.send_document(message.from_user.id, open('document.docx', 'rb'))
        bot.send_document(message.from_user.id, open('document.pdf', 'rb'))
        bot.send_message(message.from_user.id, "Обязательно проверь правильность")
        bot.send_message(message.from_user.id, "Создать новый документ /doc\n\n"
                                               "Другие команды /help")
    else:
        if re.match('\d\d', temp_date) is not None:
            for i, element in enumerate(temp_date.split('.')):
                last_date[i] = str(element)

            for i, element in enumerate(last_date):
                if len(element) == 1:
                    last_date[i] = '0' + element
            temp_date = '.'.join(last_date)
            bot.send_message(message.from_user.id, f'Укажи количество часов для даты: {temp_date}')
            bot.register_next_step_handler(message, get_hours)
        else:
            bot.send_message(message.from_user.id, 'Дата указана некорректно, попробуй ещё')
            bot.register_next_step_handler(message, get_date)


def get_hours(message):
    global data, temp_date
    temp_hours = message.text
    data.append([temp_date, temp_hours])
    bot.send_message(message.from_user.id, 'Укажи cледующую дату или нажми /end, чтобы закончить')
    temp_date = None
    bot.register_next_step_handler(message, get_date)


def create_doc():
    global data, first_str, price, sum_hours, total_sum, last_date, month

    doc = Document("Schet_na_oplatu.docx")

    table = doc.add_table(rows=len(data) + 2, cols=6)
    modifyBorder(table)
    table = set_col_widths(table)

    for i, element in enumerate(first_str):
        field = table.cell(0, i).paragraphs[0]
        field.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        field = field.add_run(element)
        field.bold = True
        field.font.size = Pt(9)

    for i, mas in enumerate(data):
        for j, element in enumerate(mas):
            if not j:
                # пишем дату
                field = table.cell(i + 1, j + 1).paragraphs[0] \
                    .add_run(f'Услуги экскаватора-погрузчика {element}')
            else:
                # пишем кол-во часов
                field = table.cell(i + 1, j + 1).paragraphs[0]
                field.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                field = field.add_run(str(element))
                field.font.size = Pt(8)
                sum_hours += int(element)

                # считаем сумму за день
                field = table.cell(i + 1, j + 4).paragraphs[0]
                field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                day_price = int(element) * price
                field = field.add_run(f'{day_price:,.2f}'.replace(',', ' '))
                total_sum += day_price
            field.font.size = Pt(8)

    # прописываем стандартные значения
    for i in range(len(data)):
        # порядковый номер строки
        field = table.cell(i + 1, 0).paragraphs[0]
        field.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        field = field.add_run(str(i + 1))
        field.font.size = Pt(8)

        # надпись час
        field = table.cell(i + 1, 3).paragraphs[0]
        field.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        field = field.add_run('час')
        field.font.size = Pt(8)

        # цена
        field = table.cell(i + 1, 4).paragraphs[0]
        field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        field = field.add_run(f'{price:,.2f}'.replace(',', ' '))
        field.font.size = Pt(8)

    field = table.cell(len(data) + 1, 2).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    field = field.add_run(str(sum_hours))
    field.font.size = Pt(8)

    field = table.cell(len(data) + 1, 5).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    field = field.add_run(f'{total_sum:,.2f}'.replace(',', ' '))
    field.font.size = Pt(8)

    doc.add_paragraph().add_run('')

    table = doc.add_table(rows=3, cols=2)
    widths = (Cm(17.5), Cm(3))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    field = table.cell(0, 0).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    field = field.add_run('Итого:')
    field.font.size = Pt(9)
    field.font.name = 'Arial'
    field.bold = True

    field = table.cell(1, 0).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    field = field.add_run('Без налога (НДС)')
    field.font.size = Pt(9)
    field.font.name = 'Arial'
    field.bold = True

    field = table.cell(2, 0).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    field = field.add_run('Всего к оплате:')
    field.font.size = Pt(9)
    field.font.name = 'Arial'
    field.bold = True

    field = table.cell(0, 1).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    field = field.add_run(f'{total_sum:,.2f}'.replace(',', ' '))
    field.font.size = Pt(9)
    field.font.name = 'Arial'
    field.bold = True

    field = table.cell(1, 1).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    field = field.add_run('–')
    field.font.size = Pt(9)
    field.font.name = 'Arial'
    field.bold = True

    field = table.cell(2, 1).paragraphs[0]
    field.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    field = field.add_run(f'{total_sum:,.2f}'.replace(',', ' '))
    field.font.size = Pt(9)
    field.font.name = 'Arial'
    field.bold = True

    # добавляем информацию после таблицы

    doc.add_paragraph().add_run('')

    paragraph = doc.add_paragraph().add_run(f'Всего наименований {len(data)}, на сумму ' +
                                            (f'{total_sum:,.2f}'.replace(',', ' ')) + ' RUB')
    paragraph.font.size = Pt(9)

    paragraph = doc.add_paragraph().add_run(f'{num2text(total_sum).capitalize()} рублей 00 копеек')
    paragraph.bold = True
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'

    doc.add_paragraph().add_run('')

    paragraph = doc.add_paragraph().add_run('ВНИМАНИЕ! Счет действителен в течение 3-х календарных дней.')
    paragraph.font.size = Pt(9)

    doc.add_paragraph("____________________________________________________________________________________")
    doc.add_paragraph().add_run('')

    paragraph = doc.add_paragraph()
    paragraph.add_run('                                     ')
    paragraph.add_run('                                     ')\
        .font.underline = True

    paragraph.add_run('          ')
    paragraph = paragraph.add_run('                           Чернов Э.В.                           ')
    paragraph.font.underline = True
    paragraph.font.size = Pt(9)

    paragraph = doc.add_paragraph()
    paragraph.add_run('                                     ')
    paragraph.add_run('                           подпись                    ')\
        .font.size = Pt(6)
    paragraph.add_run('                                                             расшифровка подписи ')\
        .font.size = Pt(6)

    doc.add_paragraph().add_run('')

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph = paragraph.add_run('Режим НО: НПД')
    paragraph.bold = True
    paragraph.font.size = Pt(9)
    field.font.name = 'Arial'

    now = datetime.datetime.now()
    doc.paragraphs[2].text = \
        f'Счет на оплату № {number} от {str(now.day) + " " +  str(month[now.month - 1]) + " " + str(now.year)} г.'

    doc.save('document.docx')

    generate_pdf("document.docx")

    sum_hours = 0
    total_sum = 0
    data = last_date = []


def set_col_widths(table):
    widths = (Cm(1.1), Cm(8.66), Cm(1.91),
              Cm(1.23), Cm(1.78), Cm(2.79))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    return table


def modifyBorder(table):
    tbl = table._tbl  # get xml element in table
    total_cell = tbl.col_count * len(data) + 12
    i = 0
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr  # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')

        if i != total_cell - 2 and i != total_cell - 3:
            if i != total_cell - 1 and i != total_cell - 5:
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '6')
                tcBorders.append(left)

            if i != total_cell - 6:
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')
                right.set(qn('w:sz'), '6')
                tcBorders.append(right)

        tcBorders.append(top)
        tcBorders.append(bottom)

        tcPr.append(tcBorders)
        i += 1


def generate_pdf(doc_path):
    import os
    os.system("lowriter --convert-to pdf" + str(" ") + str(doc_path))


bot.polling(none_stop=True)
